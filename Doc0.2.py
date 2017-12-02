from docx import Document
import cyrtranslit

print("NOTE: File koji pokusavate konvertovarti stavite u isti folder kao program.")
docf = input("Ime fajla: ")
docext = input("\nExtenzija? \n1. Doc (Word 2000-2003) \n2. DocX (Word 2007-20**) \n3. Ne znam? \nNOTE: Unesite broj koji je ispred.\nUnos: ")

extension = [".doc",".docx"]

if(int(docext) == 1):
    try:
        doc = Document(docf+".doc")
    except:
        print("Fajl ne postoji\n\n")
elif(int(docext) == 2):
    doc = Document(docf+".docx")
else:
    try:
        doc = Document(docf+".doc")
        docext =1
    except:
        doc = Document(docf+".docx")
        docext =2
        
perc10 = round(len(doc.paragraphs)/10)
count = 0
perc_multi =1

print("\nProgress | █0%       █100%")
print("Progress | █",end="")
for p in doc.paragraphs:
    inline = p.runs
    for i in range(len(inline)):
        text = inline[i].text.replace(inline[i].text, cyrtranslit.to_latin(inline[i].text))
        inline[i].text = text
    count +=1
    if(count == perc10*perc_multi):
        print("█",end="")
        perc_multi +=1     
    #print (p.text)
print("█")
print("Sacuvano kao: "+docf+"_lat"+str(extension[int(docext)-1]))
doc.save(docf+"_lat"+extension[int(docext)-1])
